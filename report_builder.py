"""DOCX 보고서 빌더 — 에디토리얼 톤 (Rimberio Impact Report 레퍼런스).

디자인 시스템:
  PAPER      #F5EFE2  크림 페이퍼 (페이지 배경)
  ESPRESSO   #2A1F15  메인 다크 (디스플레이 텍스트·강조 박스)
  CREAM      #EDE6D5  크림 액센트 (다크 박스의 텍스트, 라이트 박스 배경)
  MUTED      #6B6155  보조 텍스트
  BODY       #3A2A1E  본문 텍스트
  LINE       #3A2A1E33 얇은 디바이더 (alpha 표현은 hex로 충분)

  Display 폰트: 바탕 (Korean serif) / Cambria 영문 fallback
  Body 폰트:    맑은 고딕

레이아웃 시그니처:
  - 표지 헤로 (massive serif display)
  - 작은 영문 아이브로 ("2026  VIP Briefing Report")
  - 얇은 가로선 디바이더
  - 4-스탯 그리드 (라벨 위, 값 아래)
  - 듀얼 하이라이트 박스 (다크 에스프레소 + 크림)
  - 세리프 큰 숫자 (01, 02, 03) 상위 종목 카드
  - 다크 RM 코멘트 박스
  - 크림 NOTE 사이드 패널
"""
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─── 팔레트 ──────────────────────────────────────────────────────────────────
PAPER_HEX = "F5EFE2"
ESPRESSO_HEX = "2A1F15"
ESPRESSO_LITE_HEX = "3A2A1E"
CREAM_HEX = "EDE6D5"
MUTED_HEX = "6B6155"
BODY_HEX = "3A2A1E"

ESPRESSO = RGBColor(0x2A, 0x1F, 0x15)
CREAM = RGBColor(0xED, 0xE6, 0xD5)
MUTED = RGBColor(0x6B, 0x61, 0x55)
BODY = RGBColor(0x3A, 0x2A, 0x1E)
PAPER = RGBColor(0xF5, 0xEF, 0xE2)
RISK_RED = RGBColor(0xB9, 0x1C, 0x1C)
GAIN_GREEN = RGBColor(0x16, 0xA3, 0x4A)

FONT_SERIF = "바탕"
FONT_BODY = "맑은 고딕"

NOT_AVAILABLE = "조회 불가"


# ─── 유틸 ────────────────────────────────────────────────────────────────────
def _fp(v) -> str:
    if v in (None, NOT_AVAILABLE):
        return "-"
    try:
        return f"{int(v):,}원"
    except (TypeError, ValueError):
        return str(v)


def _fl(v) -> str:
    if v in (None, NOT_AVAILABLE):
        return "-"
    try:
        n = float(v)
        if n >= 1e12: return f"{n/1e12:.1f}조"
        if n >= 1e8:  return f"{n/1e8:.0f}억"
        return f"{n:,.0f}원"
    except (TypeError, ValueError):
        return str(v)


def _today_kr() -> str:
    d = datetime.now()
    return f"{d.year}년 {d.month}월 {d.day}일"


def _add_run(paragraph, text, *, bold=False, size=10, color=BODY, font=FONT_BODY):
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font if font == FONT_BODY else "Cambria")
    rFonts.set(qn("w:hAnsi"), font if font == FONT_BODY else "Cambria")
    return run


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_no_borders(cell):
    """셀 테두리 모두 제거 (페이퍼 위에서 무겁지 않게)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _table_no_borders(table):
    tblPr = table._tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _table_thin_borders(table, hex_color, size=4):
    tblPr = table._tbl.find(qn("w:tblPr"))
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), hex_color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _para_border_bottom(paragraph, hex_color, size=4, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), hex_color)
    pBdr.append(b)


def _para_border_top(paragraph, hex_color, size=4, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    b = OxmlElement("w:top")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), hex_color)
    pBdr.append(b)


def _set_page_background(doc, hex_color):
    """문서 전체 페이지 배경색 설정 (OOXML 직접 조작).
    Word의 '디자인 > 페이지 색'에 해당. 화면·인쇄 모두 적용되도록 settings에도 추가.
    """
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    doc.element.insert(0, bg)

    settings = doc.settings.element
    if settings.find(qn("w:displayBackgroundShape")) is None:
        display_bg = OxmlElement("w:displayBackgroundShape")
        settings.append(display_bg)


def _init_doc():
    """크림 배경 + 기본 폰트 + 여백 설정한 새 Document."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    _set_page_background(doc, PAPER_HEX)
    return doc


# ─── 빌딩 블록 ───────────────────────────────────────────────────────────────
def _hero(doc, eyebrow, title, subtitle, meta):
    """표지 헤로 — 큰 세리프 디스플레이 + 얇은 하단선 + 메타."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _add_run(p, eyebrow, bold=True, size=10, color=ESPRESSO, font=FONT_BODY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.space_after = Pt(0)
    _add_run(p2, title, bold=False, size=48, color=ESPRESSO, font=FONT_SERIF)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(20)
    _add_run(p3, subtitle, size=20, color=MUTED, font=FONT_SERIF)
    _para_border_bottom(p3, ESPRESSO_HEX, size=6, space=12)

    pm = doc.add_paragraph()
    pm.paragraph_format.space_before = Pt(8)
    pm.paragraph_format.space_after = Pt(18)
    _add_run(pm, meta, size=9, color=MUTED)


def _section_eyebrow(doc, text):
    """작은 영문/한글 아이브로 (Bold 산세리프 11pt)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, text, bold=True, size=11, color=ESPRESSO)


def _section_display(doc, text):
    """큰 세리프 섹션 타이틀 (PDF의 'Audience Engagement' 같은 스타일)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, text, size=26, color=ESPRESSO, font=FONT_SERIF)


def _body_para(doc, text, *, color=BODY, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    _add_run(p, text, size=size, color=color)


def _divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    _para_border_bottom(p, ESPRESSO_HEX, size=4, space=1)


def _stat_grid(doc, items: list[tuple[str, str, RGBColor | None]]):
    """4-스탯 행: 라벨(작은 회색) → 값(큰 볼드 세리프). items = [(label, value, val_color)]."""
    n = len(items)
    table = doc.add_table(rows=2, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_no_borders(table)
    # 라벨 행
    for j, (label, _, _) in enumerate(items):
        c = table.rows[0].cells[j]
        _cell_no_borders(c)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        _add_run(c.paragraphs[0], label, size=9, color=MUTED, bold=True)
    # 값 행
    for j, (_, val, val_color) in enumerate(items):
        c = table.rows[1].cells[j]
        _cell_no_borders(c)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        _add_run(c.paragraphs[0], val, bold=False, size=22,
                 color=(val_color or ESPRESSO), font=FONT_SERIF)
    # 상하 얇은 가로선 (스탯 그리드를 페이지에서 구분)
    _para_border_top(table.rows[0].cells[0].paragraphs[0], ESPRESSO_HEX, size=4)
    _para_border_bottom(table.rows[1].cells[0].paragraphs[0], ESPRESSO_HEX, size=4)


def _dual_highlight(doc, dark_text: str, cream_text: str):
    """다크 에스프레소 + 크림 박스가 좌우로 놓인 시그니처 박스."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_no_borders(table)
    # 다크
    c1 = table.rows[0].cells[0]
    _shade_cell(c1, ESPRESSO_HEX)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    c1.paragraphs[0].paragraph_format.space_before = Pt(16)
    c1.paragraphs[0].paragraph_format.space_after = Pt(16)
    _add_run(c1.paragraphs[0], dark_text, bold=True, size=11, color=CREAM)
    # 크림
    c2 = table.rows[0].cells[1]
    _shade_cell(c2, CREAM_HEX)
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    c2.paragraphs[0].paragraph_format.space_before = Pt(16)
    c2.paragraphs[0].paragraph_format.space_after = Pt(16)
    _add_run(c2.paragraphs[0], cream_text, bold=True, size=11, color=ESPRESSO)


def _kv_grid(doc, headers: list[str], values: list[str]):
    """5-cols 가로 KV 표 (PDF 표 스타일: 헤더 다크 + 값 페이퍼)."""
    table = doc.add_table(rows=2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_thin_borders(table, ESPRESSO_LITE_HEX + "55", size=2)
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        _shade_cell(c, ESPRESSO_HEX)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_before = Pt(6)
        c.paragraphs[0].paragraph_format.space_after = Pt(6)
        _add_run(c.paragraphs[0], h, bold=True, size=9, color=CREAM)
    for j, v in enumerate(values):
        c = table.rows[1].cells[j]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_before = Pt(8)
        c.paragraphs[0].paragraph_format.space_after = Pt(8)
        _add_run(c.paragraphs[0], v, size=10, color=ESPRESSO)


def _rank_card(doc, rank: int, name: str, ticker: str,
               kv_headers: list[str], kv_values: list[str],
               analysis_lines: list[str]):
    """세리프 큰 숫자(01/02/03) + 종목명 + KV 표 + 분석 본문."""
    # 랭크 번호
    p_rank = doc.add_paragraph()
    p_rank.paragraph_format.space_before = Pt(20)
    p_rank.paragraph_format.space_after = Pt(0)
    _add_run(p_rank, f"{rank:02d}", size=44, color=ESPRESSO, font=FONT_SERIF)
    _para_border_bottom(p_rank, ESPRESSO_HEX + "33", size=4, space=8)

    # 종목명
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(2)
    p_name.paragraph_format.space_after = Pt(10)
    _add_run(p_name, name, bold=False, size=22, color=ESPRESSO, font=FONT_SERIF)
    _add_run(p_name, f"   {ticker}", size=11, color=MUTED, font=FONT_BODY)

    # KV 표
    if kv_headers:
        _kv_grid(doc, kv_headers, kv_values)

    # 분석 본문
    for line in analysis_lines:
        _body_para(doc, line)


def _rm_box(doc, lines: list[str]):
    """RM 종합 의견 — 일반 단락 (다크 박스 제거, 본문 위에 검은 굵은 글씨)."""
    _section_eyebrow(doc, "RM Comment  ·  RM 종합 의견")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        _add_run(p, line, size=10.5, color=BODY, bold=True)


def _note_panel(doc, label: str, body_lines: list[str]):
    """크림 NOTE 사이드 패널 (PDF의 NOTE 박스 스타일)."""
    table = doc.add_table(rows=1, cols=1)
    _table_no_borders(table)
    c = table.rows[0].cells[0]
    _shade_cell(c, CREAM_HEX)

    p0 = c.paragraphs[0]
    p0.paragraph_format.space_before = Pt(14)
    p0.paragraph_format.space_after = Pt(6)
    _add_run(p0, label.upper(), bold=True, size=10, color=ESPRESSO)

    for line in body_lines:
        p = c.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.45
        _add_run(p, line, size=10, color=ESPRESSO)

    p_end = c.add_paragraph()
    p_end.paragraph_format.space_before = Pt(10)


def _holdings_table(doc, sorted_p: list[dict], stock_data: dict, full=False):
    """보유 종목 표 — 깔끔한 경계선 표 형식 (한 행 = 한 줄)."""
    if full:
        # 티커 컬럼 제거 — 한 행에 한 줄로 들어가도록 나머지 컬럼 너비 확보
        headers = ["#", "종목명", "비중", "현재가", "평가금액", "손익율", "PER", "PBR", "외국인"]
        widths_cm = [0.6, 2.4, 1.2, 1.6, 2.0, 1.6, 1.4, 1.2, 1.4]
        body_size = 8
        header_size = 9
    else:
        headers = ["#", "종목명", "비중", "현재가", "평가금액", "손익율"]
        widths_cm = [0.7, 4.0, 1.3, 1.8, 2.5, 1.5]
        body_size = 9
        header_size = 10

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    table = doc.add_table(rows=len(sorted_p) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _table_thin_borders(table, ESPRESSO_LITE_HEX + "55", size=2)

    # 모든 행에 컬럼 너비 적용
    for row in table.rows:
        for j, w in enumerate(widths_cm):
            row.cells[j].width = Cm(w)

    # 헤더 (다크 배경 + 크림 텍스트)
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        _shade_cell(c, ESPRESSO_HEX)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 1
                                     else WD_ALIGN_PARAGRAPH.CENTER)
        c.paragraphs[0].paragraph_format.space_before = Pt(4)
        c.paragraphs[0].paragraph_format.space_after = Pt(4)
        _add_run(c.paragraphs[0], h, bold=True, size=header_size, color=CREAM)

    # 바디
    for i, h in enumerate(sorted_p, 1):
        d = stock_data.get(h.get("ticker") or "", {}) or {}
        krx = d.get("krx", {}) or {}
        nf = d.get("naver_finance", {}) or {}
        price = krx.get("current_price")
        ev = h.get("평가금액")
        buy = h.get("매입가")
        pnl = h.get("손익율")
        if pnl is None and isinstance(buy, (int, float)) and isinstance(price, (int, float)) and buy:
            pnl = (price - buy) / buy * 100
        foreign = krx.get("foreign_ownership_ratio")

        pnl_color = (GAIN_GREEN if pnl is not None and pnl >= 0
                     else RISK_RED if pnl is not None else ESPRESSO)

        if full:
            cells = [
                (f"{i}", ESPRESSO, "center"),
                (h["종목명"], ESPRESSO, "left"),
                (f"{h.get('비중',0):.1f}%", ESPRESSO, "center"),
                (f"{int(price):,}" if isinstance(price, (int, float)) and price != NOT_AVAILABLE else "-", ESPRESSO, "right"),
                (f"{int(ev):,}" if isinstance(ev, (int, float)) else "-", ESPRESSO, "right"),
                (f"{pnl:+.2f}%" if pnl is not None else "-", pnl_color, "right"),
                (str(nf.get("per", "-")), ESPRESSO, "center"),
                (str(nf.get("pbr", "-")), ESPRESSO, "center"),
                (f"{foreign}%" if isinstance(foreign, (int, float)) else "-", ESPRESSO, "center"),
            ]
        else:
            cells = [
                (f"{i}", MUTED, "center"),
                (h["종목명"], ESPRESSO, "left"),
                (f"{h.get('비중',0):.1f}%", ESPRESSO, "center"),
                (f"{int(price):,}" if isinstance(price, (int, float)) and price != NOT_AVAILABLE else "-", ESPRESSO, "right"),
                (f"{int(ev):,}" if isinstance(ev, (int, float)) else "-", ESPRESSO, "right"),
                (f"{pnl:+.2f}%" if pnl is not None else "-", pnl_color, "right"),
            ]

        for j, (val, color, align) in enumerate(cells):
            c = table.rows[i].cells[j]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[0].alignment = align_map[align]
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
            c.paragraphs[0].paragraph_format.space_after = Pt(3)
            bold = (j == 1)
            _add_run(c.paragraphs[0], val, bold=bold, size=body_size,
                     color=color if not bold else ESPRESSO)


def _key_impact_bullets(doc, label: str, bullets: list[str]):
    """3개 핵심 포인트 불릿 (PDF의 'Key Impact' 스타일)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, label, bold=True, size=10, color=ESPRESSO)
    for b in bullets:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.left_indent = Cm(0.4)
        _add_run(bp, "•   ", size=11, color=ESPRESSO)
        _add_run(bp, b, size=10, color=BODY)


def _checkbox_list(doc, label: str, items: list[str]):
    """☐ 체크박스 리스트 (PDF의 Action Items 스타일)."""
    if label:
        _section_eyebrow(doc, label)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.2)
        _add_run(p, "☐   ", size=12, color=ESPRESSO)
        _add_run(p, item, size=10.5, color=BODY)


def _pull_quote(doc, text: str):
    """큰 이탤릭 세리프 풀쿼트 (잡지 스타일)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    run = _add_run(p, f"“{text}”", size=14, color=ESPRESSO, font=FONT_SERIF)
    run.italic = True
    _para_border_top(p, ESPRESSO_HEX + "44", size=4, space=8)
    _para_border_bottom(p, ESPRESSO_HEX + "44", size=4, space=8)


def _range_visual(doc, low, current, high):
    """52주 범위 텍스트 비주얼: 저점 ━━●━━━ 고점."""
    if not all(isinstance(v, (int, float)) and v not in (None, NOT_AVAILABLE)
               for v in (low, current, high)) or high == low:
        # 데이터 부족 — 단순 문자열로 폴백
        _body_para(doc,
                   f"52주 저점 {_fp(low)}    →    현재가 {_fp(current)}    →    52주 고점 {_fp(high)}")
        return

    pos = (current - low) / (high - low)
    pos = max(0.0, min(1.0, pos))
    total_width = 30
    marker_pos = round(pos * (total_width - 1))
    left_bar = "━" * marker_pos
    right_bar = "━" * (total_width - 1 - marker_pos)

    p_labels = doc.add_paragraph()
    p_labels.paragraph_format.space_before = Pt(8)
    p_labels.paragraph_format.space_after = Pt(2)
    _add_run(p_labels, f"저점 {_fp(low)}", size=9, color=MUTED)
    _add_run(p_labels, "      현재가 ", size=9, color=MUTED)
    _add_run(p_labels, _fp(current), bold=True, size=10, color=ESPRESSO)
    _add_run(p_labels, f"      고점 {_fp(high)}", size=9, color=MUTED)

    p_bar = doc.add_paragraph()
    p_bar.paragraph_format.space_after = Pt(10)
    _add_run(p_bar, left_bar, size=10, color=ESPRESSO)
    _add_run(p_bar, "●", size=14, color=ESPRESSO)
    _add_run(p_bar, right_bar, size=10, color=MUTED)


def _sector_bars(doc, sector_agg: dict[str, float]):
    """섹터 비중 가로 막대 (텍스트 차트). PDF의 파이차트를 텍스트로 변환."""
    if not sector_agg:
        return
    items = sorted(sector_agg.items(), key=lambda x: -x[1])
    max_w = max(v for _, v in items) or 1
    for sector, w in items:
        bars = round((w / max_w) * 26)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _add_run(p, sector.ljust(12, "　"), size=10, color=ESPRESSO)
        _add_run(p, "  ", size=10)
        _add_run(p, "█" * bars, size=10, color=ESPRESSO)
        _add_run(p, "░" * (26 - bars), size=10, color=MUTED)
        _add_run(p, f"  {w:.1f}%", bold=True, size=10, color=ESPRESSO)


def _generate_next_steps(portfolio, sector_agg, ai_analysis) -> list[str]:
    """Next Steps 체크리스트 — 데이터 기반 자동 생성."""
    steps = ["분기별 포트폴리오 리밸런싱 검토 일정 수립"]

    heavy = [k for k, v in sector_agg.items() if v > 30]
    if heavy:
        steps.append(f"{', '.join(heavy)} 섹터 편중 완화를 위한 비중 조정 검토")
    else:
        steps.append("현재 섹터 분산 수준을 유지하며 상위 종목 모니터링 지속")

    pnl_arr = [h for h in portfolio if h.get("손익율") is not None]
    if pnl_arr:
        losers = [h for h in pnl_arr if h["손익율"] < -10]
        if losers:
            steps.append(f"손실폭 -10% 초과 {len(losers)}개 종목에 대한 보유 의사결정 재검토")

    if ai_analysis:
        steps.append("상위 종목 AI 분석에서 제시된 단기·중기 전략을 다음 분기 운용 계획에 반영")

    steps.append("시장 변동성 모니터링 및 헤지 전략 정기 점검")
    return steps


# ─── Executive Summary 자동 생성 ─────────────────────────────────────────────
def _exec_summary_prose(portfolio: list[dict]) -> str:
    sorted_p = sorted(portfolio, key=lambda h: h.get("비중", 0), reverse=True)
    n = len(portfolio)
    total_ev = sum(h.get("평가금액") or 0 for h in portfolio)
    pnl_arr = [h for h in portfolio if h.get("손익율") is not None]
    avg_pnl = sum(h["손익율"] for h in pnl_arr) / len(pnl_arr) if pnl_arr else None
    profit_cnt = sum(1 for h in pnl_arr if h["손익율"] >= 0)
    top = sorted_p[0] if sorted_p else None

    parts = [f"본 보고서는 VIP 고객의 보유 포트폴리오를 정리한 종합 분석 자료입니다. "
             f"현재 총 {n}개 종목으로 구성되어 있으며,"]
    if total_ev > 0:
        parts.append(f" 총 평가금액은 {_fl(total_ev)} 규모입니다.")
    else:
        parts.append(" 평가금액 정보는 일부 종목에서 확인되지 않았습니다.")
    if avg_pnl is not None:
        direction = "양호한 수익 흐름" if avg_pnl >= 0 else "조정 국면"
        parts.append(f" 평균 손익율은 {avg_pnl:+.2f}%로 {direction}을 보이고 있으며, "
                     f"손익율이 집계된 {len(pnl_arr)}개 종목 중 {profit_cnt}개가 수익 구간에 진입했습니다.")
    if top:
        parts.append(f" 최대 비중 종목은 {top['종목명']} ({top.get('비중',0):.1f}%)이며,")
        parts.append(" 본 보고서의 상위 3개 종목 심층 분석에서 밸류에이션·기술적 흐름·증권사 컨센서스 대비 가격대를 다룹니다.")
    return "".join(parts)


# ─── Key Impact / Pull Quote 자동 추출 ───────────────────────────────────────
def _build_key_impacts(portfolio, sorted_p, total_ev, avg_pnl, pnl_arr, profit_cnt,
                      sector_agg: dict[str, float]) -> list[str]:
    bullets = []
    if sorted_p:
        top = sorted_p[0]
        bullets.append(f"최대 비중 종목은 {top['종목명']}로 전체 포트폴리오의 "
                       f"{top.get('비중',0):.1f}%를 차지합니다.")
    heavy = sorted([(k, v) for k, v in sector_agg.items() if v > 25],
                   key=lambda x: -x[1])
    if heavy:
        top_sectors = "·".join(k for k, _ in heavy[:2])
        bullets.append(f"{top_sectors} 중심의 섹터 구성으로, 섹터 편중 모니터링이 필요한 수준입니다.")
    else:
        sectors_top = sorted(sector_agg.items(), key=lambda x: -x[1])[:3]
        if sectors_top:
            bullets.append(f"{', '.join(s for s, _ in sectors_top)} 등 다양한 섹터에 분산 배분되어 있습니다.")
    if pnl_arr:
        rate = profit_cnt / len(pnl_arr) * 100
        verdict = ("대부분 수익 구간 유지" if rate >= 70 else
                   "수익/손실 종목이 혼재" if rate >= 30 else
                   "다수 종목이 조정 국면")
        bullets.append(f"손익율 집계 {len(pnl_arr)}개 종목 중 {profit_cnt}개({rate:.0f}%) 수익 구간 — {verdict}.")
    return bullets[:3]


def _pull_quote_from_ai(ai_item: dict) -> str:
    """AI 분석에서 풀쿼트용 핵심 한 줄 추출. competitive_position 우선."""
    if not ai_item:
        return ""
    source = (ai_item.get("competitive_position")
              or ai_item.get("valuation")
              or ai_item.get("recent_performance")
              or "")
    text = source.replace("\n", " ").strip()
    if not text:
        return ""
    import re as _re
    m = _re.search(r"[^.!?]+[.!?]", text)
    first = (m.group(0) if m else text).strip()
    if len(first) > 110:
        first = first[:108].rstrip() + "…"
    return first


def _first_sentence(text: str, max_chars: int = 220) -> str:
    """첫 완결 문장만 반환. max_chars 초과 시 그 안쪽 마지막 마침표에서 끊음."""
    text = (text or "").strip()
    if not text:
        return ""
    import re as _re
    m = _re.search(r"[^.!?]+[.!?]", text)
    first = (m.group(0) if m else text).strip()
    if len(first) <= max_chars:
        return first
    snippet = first[:max_chars]
    idx = snippet.rfind(".")
    if idx > max_chars * 0.5:
        return snippet[: idx + 1].strip()
    return snippet.rstrip() + "…"


# ─── 거시 브리핑 / 건강 진단 / 리밸런싱 렌더링 ────────────────────────────────
def _macro_indices_table(doc, indices: list[dict]) -> None:
    """주요 지수 표 (지수명·현재가·일간·YTD 4컬럼)."""
    if not indices:
        return
    table = doc.add_table(rows=len(indices) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_no_borders(table)
    headers = ["지수", "현재가", "일간", "YTD"]
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                                     else WD_ALIGN_PARAGRAPH.CENTER)
        c.paragraphs[0].paragraph_format.space_before = Pt(4)
        c.paragraphs[0].paragraph_format.space_after = Pt(6)
        _add_run(c.paragraphs[0], h, bold=True, size=9, color=ESPRESSO)
        _para_border_bottom(c.paragraphs[0], ESPRESSO_HEX, size=6)

    for i, idx in enumerate(indices, 1):
        daily_color = GAIN_GREEN if idx["daily_pct"] >= 0 else RISK_RED
        ytd_color = GAIN_GREEN if idx["ytd_pct"] >= 0 else RISK_RED
        cells = [
            (idx["name"], ESPRESSO, True, "left"),
            (f"{idx['current']:,.2f}", ESPRESSO, False, "center"),
            (f"{idx['daily_pct']:+.2f}%", daily_color, False, "center"),
            (f"{idx['ytd_pct']:+.2f}%", ytd_color, False, "center"),
        ]
        for j, (val, color, bold, align) in enumerate(cells):
            c = table.rows[i].cells[j]
            c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if align == "left"
                                         else WD_ALIGN_PARAGRAPH.CENTER)
            c.paragraphs[0].paragraph_format.space_before = Pt(4)
            c.paragraphs[0].paragraph_format.space_after = Pt(4)
            _add_run(c.paragraphs[0], val, bold=bold, size=10, color=color)
            _para_border_bottom(c.paragraphs[0], ESPRESSO_HEX + "1A", size=2)


def _macro_page(doc, macro: dict) -> None:
    """거시경제 브리핑 — 전용 페이지."""
    if not macro:
        return
    _section_eyebrow(doc, "Macro Briefing  ·  거시경제 브리핑")
    _section_display(doc, "오늘의 시장")

    if macro.get("headline"):
        _pull_quote(doc, macro["headline"])

    if macro.get("indices"):
        _section_eyebrow(doc, "Key Indices  ·  주요 지수")
        _macro_indices_table(doc, macro["indices"])

    if macro.get("macro_overview"):
        _section_eyebrow(doc, "Market Overview  ·  시장 흐름")
        _body_para(doc, macro["macro_overview"])

    if macro.get("key_drivers"):
        _key_impact_bullets(doc, "Key Drivers  ·  오늘의 핵심 동인",
                            list(macro["key_drivers"])[:4])

    if macro.get("portfolio_implication"):
        _note_panel(doc, "Portfolio Implication  ·  고객 포트폴리오 시각",
                    [macro["portfolio_implication"]])


def _health_page(doc, health: dict) -> None:
    """포트폴리오 건강 진단 — 전용 페이지."""
    if not health:
        return
    _section_eyebrow(doc, "Portfolio Health  ·  포트폴리오 건강 진단")
    _section_display(doc, "포트폴리오 진단")

    conc = health.get("concentration", {}) or {}
    sect = health.get("sector", {}) or {}
    perf = health.get("performance", {}) or {}
    alpha = health.get("alpha", {}) or {}

    def _verdict_color(verdict):
        if verdict in ("고집중", "편중"):
            return RISK_RED
        if verdict in ("다소 집중", "다소 편중"):
            return RGBColor(0xCA, 0x8A, 0x04)
        return GAIN_GREEN

    top3_color = _verdict_color(conc.get("verdict"))
    sect_color = _verdict_color(sect.get("verdict"))
    avg_ret = perf.get("avg_return", 0) or 0
    ret_color = GAIN_GREEN if avg_ret >= 0 else RISK_RED
    alpha_val = alpha.get("alpha")
    alpha_color = GAIN_GREEN if (alpha_val or 0) >= 0 else RISK_RED

    _stat_grid(doc, [
        ("상위 3종목 비중", f"{conc.get('top3_weight', 0):.1f}%", top3_color),
        (f"최대섹터 {sect.get('max_sector', '-')}",
         f"{sect.get('max_sector_weight', 0):.1f}%", sect_color),
        ("평균 손익율", f"{avg_ret:+.2f}%", ret_color),
        (f"알파 vs {alpha.get('benchmark_name', 'KOSPI')}",
         f"{alpha_val:+.2f}%p" if alpha_val is not None else "-", alpha_color),
    ])

    bullets = [
        f"집중도: {conc.get('verdict', '-')} — 상위 3종목 {conc.get('top3_weight', 0):.1f}% (HHI {conc.get('hhi', 0):.0f})",
        f"섹터 분산: {sect.get('verdict', '-')} — 최대 섹터({sect.get('max_sector', '-')}) {sect.get('max_sector_weight', 0):.1f}%",
    ]
    if perf.get("total_measured"):
        bullets.append(
            f"수익 분포: 측정 {perf['total_measured']}개 중 {perf['profit_count']}개 수익 / {perf['loss_count']}개 손실"
        )
    if alpha_val is not None:
        verdict_a = "벤치마크 상회" if alpha_val >= 0 else "벤치마크 미달"
        bullets.append(
            f"성과: 포트폴리오 {alpha.get('portfolio_ytd', 0):+.2f}% vs "
            f"{alpha.get('benchmark_name', '-')} {alpha.get('benchmark_ytd', 0):+.2f}% — {verdict_a}"
        )
    _key_impact_bullets(doc, "Diagnosis  ·  종합 평가", bullets)

    if perf.get("best") and perf.get("worst"):
        _note_panel(doc, "Top / Bottom  ·  최고·최저 종목", [
            f"최고: {perf['best'][0]}  {perf['best'][1]:+.2f}%",
            f"최저: {perf['worst'][0]}  {perf['worst'][1]:+.2f}%",
        ])


def _rebalancing_section(doc, rebalancing: list[dict]) -> None:
    """리밸런싱 액션 플랜 — AI 도출 구체 액션 카드 형태."""
    if not rebalancing:
        return
    _section_eyebrow(doc, "Rebalancing Actions  ·  리밸런싱 액션 플랜")
    _body_para(doc,
               "AI가 포트폴리오 진단 결과와 시장 환경을 바탕으로 도출한 구체적 실행 액션입니다.",
               color=MUTED, size=10)

    for action in rebalancing:
        priority = action.get("priority", "-")
        action_type = action.get("action_type", "-")
        target = action.get("target", "-")
        current = action.get("current_weight", 0)
        target_w = action.get("target_weight", 0)
        rationale = action.get("rationale", "")

        # 액션 타입별 색상
        if action_type in ("축소", "매도", "현금확보"):
            action_color = RISK_RED
        elif action_type in ("확대", "신규매수"):
            action_color = GAIN_GREEN
        else:
            action_color = ESPRESSO

        # 헤더: #우선순위  액션타입  대상
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, f"#{priority}  ", bold=True, size=11, color=MUTED)
        _add_run(p, action_type, bold=True, size=11, color=action_color)
        _add_run(p, f"   {target}", size=11, color=ESPRESSO)

        # 비중 변화
        try:
            current_f = float(current)
            target_f = float(target_w)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            _add_run(p2, f"비중  {current_f:.1f}%  →  {target_f:.1f}%",
                     size=10, color=BODY, bold=True)
        except (TypeError, ValueError):
            pass

        # 근거
        if rationale:
            p3 = doc.add_paragraph()
            p3.paragraph_format.space_after = Pt(6)
            p3.paragraph_format.left_indent = Cm(0.3)
            _add_run(p3, rationale, size=10, color=BODY)
            _para_border_bottom(p3, ESPRESSO_HEX + "1A", size=2)


# ─── 종목별 신규 섹션: PER 비교 / 30일 수급 / 1년 차트 ──────────────────────
def _per_compare_block(doc, per_compare: dict) -> None:
    """PER 현재 vs 5년 평균 vs 업종 평균 — 미니 표 + 판단 한 줄."""
    if not per_compare:
        return
    cur = per_compare.get("current_per")
    avg_5y = per_compare.get("avg_5y")
    min_5y = per_compare.get("min_5y")
    max_5y = per_compare.get("max_5y")
    sector_avg = per_compare.get("sector_avg_per")
    sector = per_compare.get("sector", "-")

    if not any([cur, avg_5y, sector_avg]):
        return

    _section_eyebrow(doc, "Valuation Compare  ·  PER 역사·업종 비교")

    headers = ["현재 PER", "5년 평균", "5년 최저", "5년 최고", f"업종({sector}) 평균"]
    values = [
        f"{cur:.2f}" if isinstance(cur, (int, float)) else "-",
        f"{avg_5y:.2f}" if avg_5y else "-",
        f"{min_5y:.2f}" if min_5y else "-",
        f"{max_5y:.2f}" if max_5y else "-",
        f"{sector_avg:.1f}" if sector_avg else "-",
    ]
    _kv_grid(doc, headers, values)

    bullets = []
    if isinstance(cur, (int, float)) and avg_5y:
        vs_h = per_compare.get("vs_history_pct", 0)
        verdict = ("역사적 저평가 구간" if vs_h < -10
                   else "역사적 고평가 구간" if vs_h > 10
                   else "역사적 평균 수준")
        bullets.append(f"5년 평균 대비 {vs_h:+.0f}% — {verdict}")
    if isinstance(cur, (int, float)) and sector_avg:
        vs_s = per_compare.get("vs_sector_pct", 0)
        if vs_s > 0:
            bullets.append(f"업종 평균 대비 {vs_s:+.0f}% 프리미엄")
        else:
            bullets.append(f"업종 평균 대비 {abs(vs_s):.0f}% 디스카운트")

    for b in bullets:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.4)
        _add_run(p, "•   ", size=10, color=ESPRESSO)
        _add_run(p, b, size=10, color=BODY)


def _flow_block(doc, flow: dict) -> None:
    """30일 외국인·기관 누적 매매 — 부호별 색상."""
    if not flow:
        return
    from flow_data import format_flow_billion

    foreign = flow.get("foreign_30d_won", 0)
    inst = flow.get("institution_30d_won", 0)

    _section_eyebrow(doc, "Trading Flow  ·  30일 외국인·기관 수급")

    f_amt, f_dir = format_flow_billion(foreign)
    i_amt, i_dir = format_flow_billion(inst)

    f_color = GAIN_GREEN if foreign > 0 else RISK_RED if foreign < 0 else MUTED
    i_color = GAIN_GREEN if inst > 0 else RISK_RED if inst < 0 else MUTED

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.left_indent = Cm(0.4)
    _add_run(p1, "외국인:  ", bold=True, size=10, color=ESPRESSO)
    _add_run(p1, f_amt, bold=True, size=10, color=f_color)
    _add_run(p1, f"  ({f_dir})", size=10, color=BODY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.left_indent = Cm(0.4)
    _add_run(p2, "기관:     ", bold=True, size=10, color=ESPRESSO)
    _add_run(p2, i_amt, bold=True, size=10, color=i_color)
    _add_run(p2, f"  ({i_dir})", size=10, color=BODY)


def _chart_block(doc, chart_png: bytes) -> None:
    """1년 주가 차트 이미지 임베드 (PNG bytes)."""
    if not chart_png:
        return
    from io import BytesIO
    from docx.shared import Inches as _Inches

    _section_eyebrow(doc, "Price Chart  ·  최근 1년 주가 추이")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    try:
        run.add_picture(BytesIO(chart_png), width=_Inches(5.8))
    except Exception:
        pass


# ─── 요약보고서 ──────────────────────────────────────────────────────────────
def build_summary(portfolio, stock_data, ai_analysis,
                  sector_map, client, rm,
                  macro=None, health=None, rebalancing=None,
                  account=None, deep_data=None):
    doc = _init_doc()
    sorted_p = sorted(portfolio, key=lambda h: h.get("비중", 0), reverse=True)
    total_ev = sum(h.get("평가금액") or 0 for h in portfolio)
    pnl_arr = [h for h in portfolio if h.get("손익율") is not None]
    avg_pnl = sum(h["손익율"] for h in pnl_arr) / len(pnl_arr) if pnl_arr else None
    profit_cnt = sum(1 for h in pnl_arr if h["손익율"] >= 0)

    sector_agg: dict[str, float] = {}
    for h in portfolio:
        s = sector_map.get(h.get("ticker") or "", "기타")
        sector_agg[s] = sector_agg.get(s, 0) + h.get("비중", 0)

    # ═══════════════ PAGE 1: 표지 + 종합 요약 ═══════════════
    _hero(doc,
          eyebrow=f"{datetime.now().year}    VIP Briefing Report",
          title="VIP Briefing",
          subtitle=f"{client}님 포트폴리오 분석",
          meta=f"작성일 {_today_kr()}     ·     담당 RM {rm}     ·     보유 종목 {len(portfolio)}개")

    _section_eyebrow(doc, "Executive Summary  ·  종합 요약")
    _body_para(doc, _exec_summary_prose(portfolio))

    # 인포그래픽: 4-스탯 그리드
    pnl_color = (GAIN_GREEN if avg_pnl is not None and avg_pnl >= 0
                 else RISK_RED if avg_pnl is not None else ESPRESSO)
    _stat_grid(doc, [
        ("총 종목 수",        f"{len(portfolio)}개", ESPRESSO),
        ("총 평가금액",       _fl(total_ev) if total_ev else "-", ESPRESSO),
        ("평균 손익율",       f"{avg_pnl:+.2f}%" if avg_pnl is not None else "-", pnl_color),
        ("최대 비중 종목",    sorted_p[0]["종목명"] if sorted_p else "-", ESPRESSO),
    ])

    # 듀얼 하이라이트 — 핵심 한 줄 메시지
    dark = (f"{len(pnl_arr)}개 측정 종목 중 {profit_cnt}개가 수익 구간 진입"
            if pnl_arr else "포트폴리오 비중 분포 분석 완료")
    cream = (f"최대 비중  {sorted_p[0]['종목명']}  ·  {sorted_p[0].get('비중',0):.1f}%"
             if sorted_p else "비중 데이터 정리 완료")
    _dual_highlight(doc, dark, cream)

    # Key Impact 불릿
    impacts = _build_key_impacts(portfolio, sorted_p, total_ev, avg_pnl, pnl_arr, profit_cnt, sector_agg)
    _key_impact_bullets(doc, "Key Impact  ·  핵심 포인트", impacts)

    # 보유 종목 표
    _section_display(doc, "Portfolio Holdings")
    _body_para(doc, f"보유 {len(portfolio)}개 종목의 비중·평가·손익율을 한눈에 정리한 표입니다.",
               color=MUTED, size=10)
    _holdings_table(doc, sorted_p, stock_data, full=False)

    # ═══════════════ PAGE 2: 상위 종목 + 섹터 + Next Steps + RM ═══════════════
    doc.add_page_break()

    _section_eyebrow(doc, "Top Holdings Deep Dive  ·  상위 3개 종목")
    _section_display(doc, "AI 심층 분석")
    _body_para(doc, "비중 상위 3종목을 Claude Sonnet이 밸류에이션·기술·투자의견 관점에서 분석했습니다.",
               color=MUTED, size=10)

    top3 = sorted_p[:3]
    for i, h in enumerate(top3, 1):
        d = stock_data.get(h.get("ticker") or "", {}) or {}
        krx = d.get("krx", {}) or {}
        nf = d.get("naver_finance", {}) or {}
        ai = next((x for x in ai_analysis if x.get("name") == h["종목명"]), None)
        target = nf.get("analyst_target_price")
        target_str = f"{int(target):,}원" if isinstance(target, (int, float)) else "-"

        # 랭크 번호 + 종목명
        p_rank = doc.add_paragraph()
        p_rank.paragraph_format.space_before = Pt(20)
        p_rank.paragraph_format.space_after = Pt(0)
        _add_run(p_rank, f"{i:02d}", size=42, color=ESPRESSO, font=FONT_SERIF)
        _para_border_bottom(p_rank, ESPRESSO_HEX + "33", size=4, space=6)

        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_before = Pt(2)
        p_name.paragraph_format.space_after = Pt(2)
        _add_run(p_name, h["종목명"], size=20, color=ESPRESSO, font=FONT_SERIF)
        _add_run(p_name, f"   {h.get('ticker') or '-'}  ·  비중 {h.get('비중',0):.1f}%",
                 size=10, color=MUTED, font=FONT_BODY)

        # Pull Quote — AI 분석 첫 문장
        quote = _pull_quote_from_ai(ai)
        if quote:
            _pull_quote(doc, quote)

        # KV 표
        _kv_grid(doc,
                 headers=["현재가", "PER", "PBR", "목표주가"],
                 values=[_fp(krx.get("current_price")), str(nf.get("per", "-")),
                         str(nf.get("pbr", "-")), target_str])

    # 리밸런싱 액션 (있을 때만)
    if rebalancing:
        _rebalancing_section(doc, rebalancing)

    # Next Steps 체크리스트
    next_steps = _generate_next_steps(portfolio, sector_agg, ai_analysis)
    _checkbox_list(doc, "Next Steps  ·  실행 액션", next_steps[:4])

    # RM 코멘트 다크 박스
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    rm_lines = ([f"[{a.get('name', '?')}]  {(a.get('strategy_short') or '').strip()}" for a in ai_analysis]
                if ai_analysis else ["포트폴리오 전반적으로 안정적인 비중 구성을 유지하고 있습니다. 분기 단위 리밸런싱을 권장합니다."])
    _rm_box(doc, rm_lines)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 상세보고서 ──────────────────────────────────────────────────────────────
def build_detail(portfolio, stock_data, ai_analysis,
                 sector_map, client, rm,
                 macro=None, health=None, rebalancing=None,
                 account=None, deep_data=None):
    doc = _init_doc()
    sorted_p = sorted(portfolio, key=lambda h: h.get("비중", 0), reverse=True)
    total_ev = sum(h.get("평가금액") or 0 for h in portfolio)
    pnl_arr = [h for h in portfolio if h.get("손익율") is not None]
    avg_pnl = sum(h["손익율"] for h in pnl_arr) / len(pnl_arr) if pnl_arr else None
    profit_cnt = sum(1 for h in pnl_arr if h["손익율"] >= 0)

    # ─── 페이지 1: 표지 + Summary ──
    _hero(doc,
          eyebrow=f"{datetime.now().year}    VIP Briefing Report  ·  Detail",
          title="VIP Briefing",
          subtitle=f"{client}님 포트폴리오 심층 분석",
          meta=f"작성일 {_today_kr()}     ·     담당 RM {rm}     ·     보유 종목 {len(portfolio)}개")

    _section_eyebrow(doc, "Executive Summary  ·  종합 요약")
    _body_para(doc, _exec_summary_prose(portfolio))

    pnl_color = (GAIN_GREEN if avg_pnl is not None and avg_pnl >= 0
                 else RISK_RED if avg_pnl is not None else ESPRESSO)
    _stat_grid(doc, [
        ("총 종목 수",        f"{len(portfolio)}개", ESPRESSO),
        ("총 평가금액",       _fl(total_ev) if total_ev else "-", ESPRESSO),
        ("평균 손익율",       f"{avg_pnl:+.2f}%" if avg_pnl is not None else "-", pnl_color),
        ("최대 비중 종목",    sorted_p[0]["종목명"] if sorted_p else "-", ESPRESSO),
    ])

    dark = (f"{len(pnl_arr)}개 측정 종목 중 {profit_cnt}개가 수익 구간 진입"
            if pnl_arr else "포트폴리오 비중 분포 분석 완료")
    cream = (f"최대 비중  {sorted_p[0]['종목명']}  ·  {sorted_p[0].get('비중',0):.1f}%"
             if sorted_p else "비중 데이터 정리 완료")
    _dual_highlight(doc, dark, cream)

    # Key Impact 불릿 — 핵심 포인트 강조
    sector_agg: dict[str, float] = {}
    for h in portfolio:
        s = sector_map.get(h.get("ticker") or "", "기타")
        sector_agg[s] = sector_agg.get(s, 0) + h.get("비중", 0)
    impacts = _build_key_impacts(portfolio, sorted_p, total_ev, avg_pnl, pnl_arr, profit_cnt, sector_agg)
    _key_impact_bullets(doc, "Key Impact  ·  핵심 포인트", impacts)

    # ═══════════════ NEW PAGE: 거시경제 브리핑 ═══════════════
    if macro:
        doc.add_page_break()
        _macro_page(doc, macro)

    # ═══════════════ PAGES 2~4: 상위 종목 매거진 스프레드 ═══════════════
    top3 = sorted_p[:3]
    for idx, h in enumerate(top3, 1):
        doc.add_page_break()
        d = stock_data.get(h.get("ticker") or "", {}) or {}
        krx = d.get("krx", {}) or {}
        nf = d.get("naver_finance", {}) or {}
        ai = next((x for x in ai_analysis if x.get("name") == h["종목명"]), None) or {}

        # 랭크 헤로
        p_rank = doc.add_paragraph()
        p_rank.paragraph_format.space_before = Pt(10)
        p_rank.paragraph_format.space_after = Pt(0)
        _add_run(p_rank, f"{idx:02d}", size=56, color=ESPRESSO, font=FONT_SERIF)

        p_name = doc.add_paragraph()
        p_name.paragraph_format.space_before = Pt(0)
        p_name.paragraph_format.space_after = Pt(4)
        _add_run(p_name, h["종목명"], size=32, color=ESPRESSO, font=FONT_SERIF)

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(10)
        _add_run(p_sub, f"{h.get('ticker') or '-'}  ·  비중 {h.get('비중',0):.1f}%",
                 size=11, color=MUTED)
        _para_border_bottom(p_sub, ESPRESSO_HEX, size=4, space=10)

        # Pull Quote — competitive_position 첫 문장
        quote = _pull_quote_from_ai(ai)
        if quote:
            _pull_quote(doc, quote)

        # 종목 개요 (5컬럼 × 2행 KV 그리드)
        _section_eyebrow(doc, "Snapshot  ·  종목 개요")
        target = nf.get("analyst_target_price")
        target_str = f"{int(target):,}원" if isinstance(target, (int, float)) else "-"
        foreign = krx.get("foreign_ownership_ratio")
        foreign_str = f"{foreign}%" if isinstance(foreign, (int, float)) else "-"

        _kv_grid(doc,
                 headers=["현재가", "PER", "PBR", "EPS", "BPS"],
                 values=[_fp(krx.get("current_price")), str(nf.get("per", "-")),
                         str(nf.get("pbr", "-")), str(nf.get("eps", "-")), str(nf.get("bps", "-"))])
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        _kv_grid(doc,
                 headers=["배당수익률", "목표주가", "시가총액", "외국인비율", "비중"],
                 values=[str(nf.get("dividend_yield", "-")), target_str,
                         _fl(krx.get("market_cap")), foreign_str, f"{h.get('비중',0):.1f}%"])

        # 심층 데이터 (PER 비교, 차트, 수급) — deep_data가 있는 경우만
        _deep = (deep_data or {}).get(h.get("ticker") or "", {}) or {}

        # PER 역사·업종 비교
        if _deep.get("per_compare"):
            _per_compare_block(doc, _deep["per_compare"])

        # 52주 범위 비주얼
        _section_eyebrow(doc, "52-Week Range  ·  52주 가격 범위")
        _range_visual(doc, krx.get("week52_low"), krx.get("current_price"), krx.get("week52_high"))

        # 1년 주가 차트 이미지
        if _deep.get("chart_png"):
            _chart_block(doc, _deep["chart_png"])

        # 30일 외국인·기관 수급
        if _deep.get("flow"):
            _flow_block(doc, _deep["flow"])

        # 산업 내 경쟁 포지션
        if ai.get("competitive_position"):
            _section_eyebrow(doc, "Competitive Position  ·  산업 내 경쟁 포지션")
            _body_para(doc, ai["competitive_position"])

        # 최근 실적과 가이던스
        if ai.get("recent_performance"):
            _section_eyebrow(doc, "Recent Performance  ·  최근 실적과 가이던스")
            _body_para(doc, ai["recent_performance"])

        # 밸류에이션 분석
        if ai.get("valuation"):
            _section_eyebrow(doc, "Investment Analysis  ·  밸류에이션 분석")
            _body_para(doc, ai["valuation"])

        # 향후 모멘텀 동인
        if ai.get("catalysts"):
            _section_eyebrow(doc, "Catalysts  ·  향후 모멘텀 동인")
            _body_para(doc, ai["catalysts"])

        # 투자 전략 — 단기 / 중기 (NOTE 패널)
        _note_panel(doc, "Investment Strategy  ·  투자 전략",
                    [f"단기:  {(ai.get('strategy_short') or '모니터링 유지').strip()}",
                     f"중기:  {(ai.get('strategy_mid') or '비중 유지 검토').strip()}"])

    # ═══════════════ 마지막 페이지: 전체 종목 + 섹터 + Next Steps + RM ═══════════════
    doc.add_page_break()

    _section_eyebrow(doc, "Full Holdings  ·  전체 종목 상세 현황")
    _holdings_table(doc, sorted_p, stock_data, full=True)

    # 리밸런싱 액션 플랜 — AI 제공 구체 액션 (없으면 기존 섹터 NOTE fallback)
    if rebalancing:
        _rebalancing_section(doc, rebalancing)
    else:
        heavy = [f"{k}({v:.1f}%)" for k, v in sector_agg.items() if v > 30]
        if heavy:
            _note_panel(doc, "Rebalancing Note  ·  섹터 편중 점검",
                        [f"집중 섹터: {', '.join(heavy)}",
                         "섹터 분산을 통한 리밸런싱 검토를 권장합니다."])
        else:
            _note_panel(doc, "Sector Distribution  ·  섹터 분산",
                        ["섹터 분산이 양호한 수준입니다.",
                         "현재 비중을 유지하면서 정기적 모니터링을 권장합니다."])

    # Next Steps 체크리스트
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    next_steps = _generate_next_steps(portfolio, sector_agg, ai_analysis)
    _checkbox_list(doc, "Next Steps  ·  실행 액션", next_steps)

    # RM 코멘트
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    rm_lines = ([f"[{a.get('name', '?')}]  {(a.get('strategy_short') or '').strip()}" for a in ai_analysis]
                if ai_analysis else ["포트폴리오 전반적으로 양호합니다. 분기 단위 리밸런싱과 정기 점검을 권장합니다."])
    _rm_box(doc, rm_lines)

    # 면책
    _body_para(doc, "【면책문구】 본 자료는 유진투자증권 서울WM센터에서 투자 참고 목적으로 작성하였으며, "
                    "특정 종목에 대한 투자 권유가 아닙니다. 주가 및 재무 데이터는 실시간 변동될 수 있으며, "
                    "AI 분석은 자동 생성된 의견으로 전문가 투자 조언을 대체하지 않습니다. "
                    "투자에 따른 모든 책임은 투자자 본인에게 귀속됩니다. 데이터 출처: KRX · 네이버금융 · 에프엔가이드",
              color=MUTED, size=8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
